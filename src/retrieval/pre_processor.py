import os
import re
from pathlib import Path
from typing import List

from docx import Document

from src.lib.logger import logger
from src.schemas.models.document import BaseDocument, MarkdownDocument, Paragraph


def _normalize_whitespace(text: str) -> str:
    """
    Normalize whitespace in text by replacing non-breaking spaces and tabs with regular spaces,
    and reducing multiple consecutive spaces to a single space.

    Args:
        text (str): The input text to normalize.

    Returns:
        str: Text with normalized whitespace.
    """
    # Replace non-breaking spaces and tabs with regular spaces
    text = text.replace("\u00a0", " ").replace("\t", " ")
    # Replace multiple consecutive spaces with a single space
    return re.sub(r"\s+", " ", text)


def _clean_heading(text: str) -> str:
    """
    Clean a heading by removing leading numbers and spaces.

    Args:
        text (str): The heading text.

    Returns:
        str: Cleaned heading text in lowercase.
    """
    return re.sub(r"^\d+(\.\d+)*\s*", "", text).strip().lower()


def _is_excluded_heading(text: str, exclude_sections: list[str]) -> bool:
    """
    Check if a heading should be excluded based on predefined sections.

    Args:
        text (str): The heading text to check
        exclude_sections (list[str]): List of section names to check against

    Returns:
        bool: True if the heading matches one of the excluded sections, False otherwise
    """
    cleaned = _clean_heading(text)
    return any(cleaned.startswith(section.lower()) for section in exclude_sections)


def _extract_paragraphs(doc: Document) -> List[Paragraph]:
    """
    Extract paragraphs and tables from the document in their natural order.

    Args:
        doc (Document): The Word document to extract content from.

    Returns:
        List[Paragraph]: List of paragraphs where each paragraph is a Pydantic model containing:
            - text (str): The paragraph text content
            - style (str): The paragraph style name
            - level (Optional[int]): The heading level if applicable
    """
    paragraphs: List[Paragraph] = []
    inside_contents_section = False

    for para in doc.paragraphs:
        text = _normalize_whitespace(para.text.strip())

        # Handle Contents section
        if text.lower() == "contents":
            inside_contents_section = True
            continue
        if inside_contents_section and para.style.name.startswith("Heading"):
            inside_contents_section = False
        if inside_contents_section:
            continue

        # Get paragraph style and level
        style = para.style.name
        level = (
            int(style[-1])
            if style.startswith("Heading") and style[-1].isdigit()
            else None
        )

        # Check for table after this paragraph
        if hasattr(para._element, "getnext") and para._element.getnext() is not None:
            next_elem = para._element.getnext()
            if next_elem.tag.endswith("tbl"):
                # Convert table to markdown directly using the table element
                markdown_table = []
                for row in next_elem.tr_lst:  # Access table rows
                    cells = []
                    for cell in row.tc_lst:  # Access cells in each row
                        # Get text content from cell
                        cell_text = "".join(
                            t.text.strip() for t in cell.xpath(".//w:t")
                        )
                        cells.append(cell_text.replace("\n", " "))
                    markdown_table.append("| " + " | ".join(cells) + " |")
                    if len(markdown_table) == 1:  # After header row
                        markdown_table.append(
                            "|" + "|".join(["---" for _ in cells]) + "|"
                        )

                # Add paragraph if not empty
                if text:
                    paragraphs.append(Paragraph(text=text, style=style, level=level))
                # Add table
                paragraphs.append(
                    Paragraph(text="\n".join(markdown_table), style="Table", level=None)
                )
                continue

        # Add regular paragraph if not empty
        if text:
            paragraphs.append(Paragraph(text=text, style=style, level=level))

    return paragraphs


def _extract_filtered_content(
    paragraphs: List[Paragraph], exclude_sections: list[str]
) -> List[str]:
    """
    Extracts main content with tables in their natural position.

    Args:
        paragraphs (List[Paragraph]): List of paragraph dictionaries
        exclude_sections (list[str]): List of section names to exclude

    Returns:
        List[str]: List of markdown-formatted content lines
    """
    filtered_content = []
    exclude_section = False
    found_first_heading = False

    for para in paragraphs:
        text = para.text
        is_heading = para.style.startswith("Heading")

        # Mark first heading found
        if is_heading:
            found_first_heading = True

        # Skip content until we find first heading
        if not found_first_heading:
            continue

        if is_heading and _is_excluded_heading(text, exclude_sections):
            exclude_section = True
            continue

        if is_heading and not _is_excluded_heading(text, exclude_sections):
            exclude_section = False

        if not exclude_section:
            if is_heading:
                level = para.level or 2
                filtered_content.append(f"{'#' * level} {text}")
            else:
                filtered_content.append(text)

    return filtered_content


def _extract_excluded_content(
    paragraphs: List[Paragraph], exclude_sections: list[str]
) -> List[str]:
    """
    Extracts content from excluded sections, maintaining proper heading formatting.
    Captures all excluded content from the start of the document.

    Args:
        paragraphs (List[Paragraph]): List of paragraph dictionaries
        exclude_sections (list[str]): List of section names to exclude

    Returns:
        List[str]: List of markdown-formatted content lines from excluded sections
    """
    excluded_content = []
    exclude_section = True
    current_section = []
    found_first_heading = False

    for para in paragraphs:
        text = para.text
        is_heading = para.style.startswith("Heading")

        # Handle first heading case
        if is_heading and not found_first_heading:
            found_first_heading = True
            # If first heading is not excluded, add previous content and stop excluding
            if not _is_excluded_heading(text, exclude_sections):
                if current_section:
                    excluded_content.extend(current_section)
                    current_section = []
                exclude_section = False
                continue

        # Handle excluded section start
        if is_heading and _is_excluded_heading(text, exclude_sections):
            # Add previous section if exists
            if current_section:
                excluded_content.extend(current_section)
            exclude_section = True
            # Format heading with proper level
            level = para.level or 2
            current_section = [f"{'#' * level} {text}"]
            continue

        # Handle non-excluded heading
        if is_heading and not _is_excluded_heading(text, exclude_sections):
            if exclude_section and current_section:
                excluded_content.extend(current_section)
                current_section = []
            exclude_section = False
            continue

        # Collect content while in excluded section
        if exclude_section:
            if para.style == "Table":
                current_section.append(text)  # Tables are already formatted
            else:
                current_section.append(text)

    # Add any remaining excluded content
    if current_section:
        excluded_content.extend(current_section)

    return excluded_content


def _extract_toc(paragraphs: List[Paragraph], exclude_sections: list[str]) -> List[str]:
    """
    Extracts the table of contents based on non-excluded headings.

    Args:
        paragraphs (List[Paragraph]): List of paragraph models containing:
            - text (str): The paragraph text content
            - style (str): The paragraph style name
            - level (Optional[int]): The heading level if applicable
        exclude_sections (list[str]): List of section names to exclude from TOC

    Returns:
        List[str]: List of markdown-formatted table of contents entries
    """
    toc_content = []
    for para in paragraphs:
        if para.style.startswith("Heading") and not _is_excluded_heading(
            para.text, exclude_sections
        ):
            level = para.level or 2
            toc_content.append(f"{'  ' * (level - 1)}- {para.text}")

    return toc_content


def _save_markdown_files(
    main_content: List[str],
    excluded_content: List[str],
    toc_content: List[str],
    output_folder: str,
) -> None:
    """
    Save extracted content into separate Markdown files.

    Args:
        main_content (List[str]): List of markdown-formatted main content lines
        excluded_content (List[str]): List of markdown-formatted excluded content lines
        toc_content (List[str]): List of markdown-formatted table of contents lines
        output_folder (str): Directory path where markdown files will be saved

    Note:
        Creates three files in the output directory:
        - 24501-filtered.md: Contains the main content
        - 24501-excluded.md: Contains the excluded sections
        - 24501-toc.md: Contains the table of contents
    """
    os.makedirs(output_folder, exist_ok=True)

    files = {
        "24501-filtered.md": main_content,
        "24501-excluded.md": excluded_content,
        "24501-toc.md": toc_content,
    }

    for filename, content in files.items():
        try:
            file_path = os.path.join(output_folder, filename)
            with open(file_path, "w", encoding="utf-8") as f:
                f.write("\n\n".join(content))
            logger.info(f"Saved {filename} to {output_folder}")
        except Exception as e:
            logger.error(f"Error saving {filename}: {e}")


def _get_doc_name_and_version(file_path: Path) -> BaseDocument:
    """
    Extract document metadata from a 3GPP specification file path.

    The 3GPP specification files follow a strict naming convention where:
    1. The spec number is written without a dot after the second digit
       (e.g., "24501" for spec "24.501")
    2. Followed by a hyphen
    3. Followed by exactly three characters representing the version in base-36
       (where digits 0-9 represent themselves and letters A-Z represent 10-35)

    For example:
        "24501-j11.docx" -> DocumentBase(
            spec="24.501",
            version="19.1.1",
            release="19"
        )
        Where:
        - "24501" becomes "24.501" (inserting dot after second digit)
        - "j11" becomes "19.1.1" (j=19 in base-36, 1=1, 1=1)
        - First component (19) becomes the release number

    Args:
        file_path (Path): The path to the 3GPP specification file.

    Returns:
        DocumentBase: A Pydantic model containing:
            - spec (str): The specification number with a dot after the second digit
            - version (str): The version string with each base-36 digit converted to decimal
                           and joined with dots
            - release (str): The release number (first component of the version)

    Raises:
        ValueError: If the filename doesn't meet the specification naming convention.
    """
    stem = file_path.stem  # e.g. "24501-j11"
    try:
        spec_part, ver_part = stem.split("-")
    except ValueError:
        raise ValueError(
            f"Filename '{stem}' does not meet specification naming convention."
        )

    # Validate spec_part length (should be at least 3 digits to insert a dot after the second digit)
    if len(spec_part) < 3:
        raise ValueError(
            f"Filename '{stem}' does not meet specification naming convention."
        )

    # Insert a decimal point after the second digit and convert to float.
    spec = spec_part[:2] + "." + spec_part[2:]

    # Process version part: convert each character using base-36
    if len(ver_part) != 3:
        raise ValueError(
            f"Filename '{stem}' does not meet specification naming convention."
        )

    version_components = []
    for char in ver_part:
        try:
            # Convert each character from base-36 to decimal
            value = int(char, 36)
            version_components.append(str(value))
        except ValueError:
            raise ValueError(f"Invalid character in version part: '{char}'")

    release = version_components[0]
    version_str = ".".join(version_components)

    document_data = BaseDocument(
        spec=spec,
        version=version_str,
        release=release,
    )

    return document_data


def docx_to_markdown(
    file_path: Path, save_markdown: bool = False, output_folder: str = "data/markdown"
) -> MarkdownDocument:
    """
    Main function to process DOCX file and extract structured content.

    Args:
        file_path (str): Path to the DOCX file.
        save_markdown (bool): Whether to save the output.
        output_folder (str): Directory to save Markdown files.

    Returns:
        MarkdownOutput: A dictionary with:
            - content (str): Filtered markdown content
            - toc (str): Table of contents in markdown format
            - doc_name (str): Name/title of the document
    """
    try:
        logger.info(f"Processing file: {file_path}")
        doc = Document(file_path)

        # Get doc spec, version and release
        doc_metadata = _get_doc_name_and_version(file_path)

        paragraphs = _extract_paragraphs(doc)

        exclude_sections = {
            "annex",
            "void",
            "foreword",
            "scope",
            "references",
            "definitions and abbreviations",
            "definitions",
            "abbreviations",
        }

        filtered_content = _extract_filtered_content(paragraphs, exclude_sections)
        excluded_content = _extract_excluded_content(paragraphs, exclude_sections)
        toc_content = _extract_toc(
            paragraphs=paragraphs, exclude_sections=exclude_sections
        )

        if save_markdown:
            _save_markdown_files(
                filtered_content, excluded_content, toc_content, output_folder
            )

        return MarkdownDocument(
            content="\n\n".join(filtered_content),
            toc="\n\n".join(toc_content),
            spec=doc_metadata.spec,
            version=doc_metadata.version,
            release=doc_metadata.release,
        )

    except Exception as e:
        logger.error(f"Error processing file {file_path}: {e}")
        raise ValueError(f"Failed to convert file: {file_path} to Markdown.")


# Example usage
if __name__ == "__main__":
    file_path = Path("data/raw/24501-j11.docx")
    result = docx_to_markdown(file_path, save_markdown=True)
    print("Content length:", len(result.content))
    print("Table of contents length:", len(result.toc))
